"""Generate Homework7 submission Word document with detailed answers."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)

# ---- Title ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BADM 566 Application of AI in Risk Management')
run.bold = True
run.font.size = Pt(14)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Homework 7')
run.bold = True
run.font.size = Pt(14)

due = doc.add_paragraph()
due.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = due.add_run('Due: 04/19/2026')
run.font.size = Pt(12)

# Name / Score
info = doc.add_paragraph()
info.add_run('Name: ').bold = True
info.add_run('Rajesh Kumar Sahoo')
info.add_run('                                        Score: _________________')

doc.add_paragraph()

# ---- Background ----
doc.add_paragraph().add_run('AI in Technology Risk Assessment').bold = True

doc.add_paragraph(
    'Technology risk is any threat to your business data, critical systems and business processes. '
    'Technology risks include hardware and software failures, human error, spam, viruses and malicious '
    'attacks, as well as natural disasters such as fires, cyclones or floods.'
)
doc.add_paragraph(
    'Technology risk assessment helps us identify vulnerabilities that put your information assets and '
    'business continuity at risk. AI makes it easy to understand current risk profile, prioritize our '
    'response, know what steps to take to remediate each issue, fine-tune risk levels to particular '
    'environment and then review the new risk profile to assess the success efforts made.'
)
doc.add_paragraph(
    'Technology risk assessment involves four key components: Threat, Vulnerability, Impact, and Likelihood. '
    'Threat refers to the frequency of an event called "threat frequency," or how often an adverse event is '
    'expected to occur. Implementing controls is an essential task to mitigate the inherent technology risk '
    'level. Even after introducing the controls in the risk processes, we might face the leftover risk called '
    'the Residual Risk.'
)

# ============================================================
# Q1
# ============================================================
doc.add_paragraph()
q1 = doc.add_paragraph()
run = q1.add_run(
    'Q1. Perform Technology Risk assessment using data in excel sheet "week7-data" in the attached Excel file '
    '"Week7-Risk_Assessment". Use column "emails" which has unstructured data, and extract topics using the '
    'sheet "Week7-Risk-Matrix". Use \'Topic Extract\' model available in MeaningCloud Add-on.'
)
run.bold = True

# ---- Q1a ----
doc.add_paragraph()
q1a = doc.add_paragraph()
run = q1a.add_run('a. Perform Risk Assessment and calculate the risk score:')
run.bold = True

doc.add_paragraph('    - Extract Technology events from the unstructured data')
doc.add_paragraph('    - Count events for each category')
doc.add_paragraph('    - Calculate likelihood and impact')
doc.add_paragraph('    - Calculate the risk levels and ratings in the attached sheet')

doc.add_paragraph()
ans_header = doc.add_paragraph()
run = ans_header.add_run('Answer:')
run.bold = True
run.underline = True

# Step 1: Topic Extraction
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 1: Extract Technology Events from Unstructured Email Data').bold = True

doc.add_paragraph(
    'Using keyword-based NLP topic extraction (as an alternative to MeaningCloud, see Note at the end), '
    'each of the 20 emails in the "week7-data" sheet was analyzed against the risk taxonomy defined in the '
    '"Week7-Risk Matrix" sheet. The following events were extracted:'
)

# Extraction table
extraction_data = [
    ('Email 1', 'Software failure', 'Testing'),
    ('Email 2', 'Software failure', 'Testing'),
    ('Email 3', 'Software failure', 'Testing, Design'),
    ('Email 4', 'Software failure', 'Testing'),
    ('Email 5', 'Software failure', 'Testing, Vulnerability'),
    ('Email 6', 'Software failure', 'Testing'),
    ('Email 7', 'Software failure', 'End of Life (EOL)'),
    ('Email 8', 'Software failure', 'End of Vendor Support (EOVS)'),
    ('Email 9', 'Hardware Failure', 'Loss of Power / Insufficient Backup'),
    ('Email 10', 'Hardware Failure', 'Hard drive failures'),
    ('Email 11', 'Hardware Failure', 'Hard drive failures'),
    ('Email 12', 'Hardware Failure', 'Natural/Man-made disasters (Fire)'),
    ('Email 13', 'Users & computers', 'Passwords never expire'),
    ('Email 14', 'Users & computers', 'Passwords not required to expire'),
    ('Email 15', 'Users & computers', 'Inactive User accounts'),
    ('Email 16', 'Permissions', 'Administrative permissions'),
    ('Email 17', 'Permissions', 'Empty security groups'),
    ('Email 18', 'Permissions', 'Administrative Groups'),
    ('Email 19', 'Data', 'Shared folders accessible by everyone'),
    ('Email 20', 'Data', 'File names containing sensitive data'),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, h in enumerate(['Email', 'Event Category', 'Sub-Event Extracted']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

for email, event, sub in extraction_data:
    row = table.add_row().cells
    row[0].text = email
    row[1].text = event
    row[2].text = sub
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# Step 2: Count events
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 2: Count Events for Each Category').bold = True

count_data = [
    ('Software failure', '10', 'Testing (6), Design (1), EOL (1), EOVS (1), Vulnerability (1), Hacking/Malware (0)'),
    ('Hardware Failure', '4', 'Hard drive failures (2), Natural/Man-made disasters (1), Loss of Power (1), Network failures (0)'),
    ('Users & computers', '3', 'Passwords never expire (1), Passwords not required to expire (1), Inactive accounts (1)'),
    ('Permissions', '3', 'Administrative permissions (1), Empty security groups (1), Administrative Groups (1)'),
    ('Data', '2', 'Shared folders accessible by everyone (1), File names with sensitive data (1)'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table2.rows[0].cells
for i, h in enumerate(['Event Category', 'Total Count', 'Sub-Event Breakdown']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

for event, count, breakdown in count_data:
    row = table2.add_row().cells
    row[0].text = event
    row[1].text = count
    row[2].text = breakdown
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# Step 3: Calculate Likelihood, Impact, Risk Factor
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 3: Calculate Likelihood, Impact, and Risk Scores').bold = True

doc.add_paragraph(
    'Following the formulas from the Risk Matrix sheet:'
)
doc.add_paragraph('    - Probability: Extracted from email text (e.g., "70% chance" = Likely = 4). '
                  'If no probability is found, default to "Likely" (rating 4) per the instructions.')
doc.add_paragraph('    - Likelihood (L): Equal to the probability rating (1-5 scale).')
doc.add_paragraph('    - Impact (I): Impact from category table x (Sub-event count / Total event type count). '
                  'For example, Testing: Impact = 4 x (6/10) = 2.40')
doc.add_paragraph('    - Risk Factor = Likelihood (L) x Impact (I)')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Risk Assessment Results:').bold = True

# Risk assessment table
risk_data = [
    ('Software failure', 'Testing', '6', '4', 'Yes', '4', '2.40', '9.60'),
    ('Software failure', 'Design', '1', '4', 'Yes', '4', '0.30', '1.20'),
    ('Software failure', 'End of Life (EOL)', '1', '4', 'No', '4', '0.30', '1.20'),
    ('Software failure', 'End of Vendor Support (EOVS)', '1', '4', 'Yes', '4', '0.20', '0.80'),
    ('Software failure', 'Vulnerability', '1', '4', 'Yes', '4', '0.30', '1.20'),
    ('Software failure', 'Hacking/Malware', '1', '4', 'Yes', '4', '0.40', '1.60'),
    ('Hardware Failure', 'Natural/Man-made disasters', '1', '4', 'No', '4', '1.25', '5.00'),
    ('Hardware Failure', 'Hard drive failures', '2', '4', 'No', '4', '2.50', '10.00'),
    ('Hardware Failure', 'Network failures', '1', '4', 'Yes', '4', '1.25', '5.00'),
    ('Hardware Failure', 'Loss of Power (Insufficient Backup)', '1', '4', 'No', '4', '1.25', '5.00'),
    ('Users & computers', 'Passwords never expire', '1', '4', 'No', '4', '1.00', '4.00'),
    ('Users & computers', 'Passwords not required to expire', '1', '4', 'Yes', '4', '0.67', '2.67'),
    ('Users & computers', 'Inactive User accounts', '1', '4', 'No', '4', '1.67', '6.67'),
    ('Permissions', 'Administrative permissions', '1', '4', 'No', '4', '1.67', '6.67'),
    ('Permissions', 'Empty security groups', '1', '4', 'No', '4', '1.67', '6.67'),
    ('Permissions', 'Administrative Groups', '1', '4', 'No', '4', '1.67', '6.67'),
    ('Data', 'Shared folders accessible by everyone', '1', '4', 'No', '4', '2.50', '10.00'),
    ('Data', 'File names containing sensitive data', '1', '4', 'No', '4', '2.50', '10.00'),
]

headers = ['Event', 'Sub-Event', 'Count', 'Prob', 'Risk Control', 'L', 'I', 'Risk (LxI)']
table3 = doc.add_table(rows=1, cols=8)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table3.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

for row_data in risk_data:
    row = table3.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    'Key Findings: The highest risk factors are Hard drive failures (10.00), Shared folders accessible by '
    'everyone (10.00), File names containing sensitive data (10.00), and Software Testing failures (9.60). '
    'These represent the most critical areas requiring immediate risk mitigation.'
)

# ============================================================
# Q1b - Linear Regression
# ============================================================
doc.add_paragraph()
q1b = doc.add_paragraph()
run = q1b.add_run(
    'b. Using Linear Regression, find the coefficients and equation of Regression after normalizing '
    'the data as in the sample regression sheet. Use Risk factor as dependent variable and features '
    'as independent variables. Convert all textual categorical variables to numeric, and clean data if necessary.'
)
run.bold = True

doc.add_paragraph()
ans_header = doc.add_paragraph()
run = ans_header.add_run('Answer:')
run.bold = True
run.underline = True

# Step 1: Data Preparation
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 1: Convert Categorical Variables to Numeric').bold = True

doc.add_paragraph(
    'Following the sample regression sheet format, all textual categorical variables were converted to numeric IDs:'
)

doc.add_paragraph('    - Event: Software failure=1, Hardware Failure=2, Users & computers=3, Permissions=4, Data=5')
doc.add_paragraph('    - Sub-Event: Sequential IDs 1 through 18')
doc.add_paragraph('    - Risk Control Present: Yes=1, No=0')

# Step 2: Normalization
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 2: Normalize Data Using Min-Max Scaling').bold = True

doc.add_paragraph(
    'All independent features were normalized to the [0, 1] range using Min-Max normalization: '
    'X_normalized = (X - X_min) / (X_max - X_min). '
    'The dependent variable (Risk Factor) was kept in its original scale.'
)

# Step 3: Regression Results
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 3: Linear Regression Results').bold = True

doc.add_paragraph(
    'A multiple linear regression was fitted with Risk Factor (L x I) as the dependent variable '
    'and the following 7 independent variables: Event ID, Sub-Event ID, Sub-Events\' Counts, '
    'Probability, Risk Control Present, Likelihood (L), and Impact (I).'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Regression Coefficients:').bold = True

coeff_data = [
    ('Intercept', '0.8018'),
    ('Event ID', '0.0042'),
    ('Sub-Event ID', '-0.0240'),
    ("Sub-Events' Counts", '-0.0165'),
    ('Probability', '~0 (negligible)'),
    ('Risk Control Present', '0.0018'),
    ('Likelihood (L)', '0 (constant in data)'),
    ('Impact (I)', '9.2129'),
]

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table4.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Coefficient'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)

for feat, coef in coeff_data:
    row = table4.add_row().cells
    row[0].text = feat
    row[1].text = coef
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Regression Equation (on normalized features):').bold = True

doc.add_paragraph(
    'Risk Factor = 0.8018 + (0.0042 x Event ID) + (-0.0240 x Sub-Event ID) + '
    '(-0.0165 x Sub-Events\' Counts) + (0.0018 x Risk Control) + (9.2129 x Impact)'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Model Performance:').bold = True

doc.add_paragraph('    - R-squared (R\u00b2) = 0.999999')
doc.add_paragraph('    - The near-perfect R\u00b2 indicates the model explains virtually all variance in risk factors.')
doc.add_paragraph(
    '    - The dominant coefficient is Impact (I) at 9.2129, confirming that Impact is the strongest '
    'predictor of the Risk Factor. This makes sense because Risk Factor = Likelihood x Impact, '
    'and Likelihood is constant at 4 across all rows (default probability rating).'
)

# ============================================================
# Q2 - SVM
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
q2 = doc.add_paragraph()
run = q2.add_run(
    'Q2. Using Support Vector Machines (SVM), predict whether the risk factor is relevant or not based on '
    'the key features like Sub-Event, Risk Control Present, Likelihood, and Impact. Please prepare the data '
    'set needed for the modeling in sheets "TechTrain" and "TechTest". Use Risk factor as dependent variable '
    'and features as independent variables. Convert all textual categorical variables to numeric, and clean '
    'data if necessary.'
)
run.bold = True

doc.add_paragraph()
ans_header = doc.add_paragraph()
run = ans_header.add_run('Answer:')
run.bold = True
run.underline = True

# Step 1: Data Preparation
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 1: Data Preparation').bold = True

doc.add_paragraph(
    'The Risk Factor was converted into a binary classification target to determine relevance:'
)
doc.add_paragraph('    - Relevant (1): Risk Factor >= 5.0 (median threshold)')
doc.add_paragraph('    - Not Relevant (0): Risk Factor < 5.0')
doc.add_paragraph(
    'This threshold was chosen based on the median risk factor value (5.0), which provides a '
    'balanced split between high-risk (relevant) and low-risk (not relevant) categories.'
)

doc.add_paragraph()
doc.add_paragraph('Categorical variables were converted to numeric:')
doc.add_paragraph('    - Sub-Event: Sequential numeric IDs (1-18), then normalized to [0, 1]')
doc.add_paragraph('    - Risk Control Present: Yes=1, No=0')
doc.add_paragraph('    - Likelihood and Impact: Already numeric, normalized using Min-Max scaling')

# Step 2: Train/Test Split
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 2: Train/Test Split').bold = True

doc.add_paragraph(
    'The 18-row dataset was split into 70% training (12 rows) and 30% testing (6 rows) '
    'using random sampling with seed=42 for reproducibility.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Training Data (TechTrain sheet - 12 records):').bold = True

train_data = [
    ('Software failure', 'Testing', '1', '4', '2.40', '1 (Relevant)'),
    ('Software failure', 'Design', '1', '4', '0.30', '0 (Not Relevant)'),
    ('Hardware Failure', 'Network failures', '1', '4', '1.25', '1 (Relevant)'),
    ('Software failure', 'Hacking/Malware', '1', '4', '0.40', '0 (Not Relevant)'),
    ('Software failure', 'End of Vendor Support', '1', '4', '0.20', '0 (Not Relevant)'),
    ('Permissions', 'Administrative permissions', '0', '4', '1.67', '1 (Relevant)'),
    ('Data', 'Shared folders accessible', '0', '4', '2.50', '1 (Relevant)'),
    ('Permissions', 'Administrative Groups', '0', '4', '1.67', '1 (Relevant)'),
    ('Users & computers', 'Passwords not req. to expire', '1', '4', '0.67', '0 (Not Relevant)'),
    ('Software failure', 'End of Life (EOL)', '0', '4', '0.30', '0 (Not Relevant)'),
    ('Hardware Failure', 'Loss of Power', '0', '4', '1.25', '1 (Relevant)'),
    ('Data', 'Sensitive data files', '0', '4', '2.50', '1 (Relevant)'),
]

table5 = doc.add_table(rows=1, cols=6)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table5.rows[0].cells
for i, h in enumerate(['Event', 'Sub-Event', 'Risk Ctrl', 'L', 'I', 'Relevant']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

for row_data in train_data:
    row = table5.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Test Data (TechTest sheet - 6 records):').bold = True

test_data = [
    ('Software failure', 'Vulnerability', '1', '4', '0.30', '0', '0'),
    ('Users & computers', 'Inactive User accounts', '0', '4', '1.67', '1', '1'),
    ('Hardware Failure', 'Hard drive failures', '0', '4', '2.50', '1', '1'),
    ('Users & computers', 'Passwords never expire', '0', '4', '1.00', '0', '1'),
    ('Permissions', 'Empty security groups', '0', '4', '1.67', '1', '1'),
    ('Hardware Failure', 'Natural/Man-made disasters', '0', '4', '1.25', '1', '1'),
]

table6 = doc.add_table(rows=1, cols=7)
table6.style = 'Table Grid'
table6.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table6.rows[0].cells
for i, h in enumerate(['Event', 'Sub-Event', 'Risk Ctrl', 'L', 'I', 'Actual', 'Predicted']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

for row_data in test_data:
    row = table6.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

# Step 3: SVM Model
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Step 3: SVM Model Training and Results').bold = True

doc.add_paragraph(
    'An SVM classifier was trained with the following configuration:'
)
doc.add_paragraph('    - Kernel: RBF (Radial Basis Function)')
doc.add_paragraph('    - Regularization parameter C: 1.0')
doc.add_paragraph('    - Gamma: scale (automatic)')
doc.add_paragraph('    - Features: Sub-Event ID, Risk Control Present, Likelihood, Impact (all normalized)')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Classification Results:').bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Confusion Matrix (Test Set):').bold = True

cm_table = doc.add_table(rows=3, cols=3)
cm_table.style = 'Table Grid'
cm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cm_table.rows[0].cells[0].text = ''
cm_table.rows[0].cells[1].text = 'Pred: Not Relevant'
cm_table.rows[0].cells[2].text = 'Pred: Relevant'
cm_table.rows[1].cells[0].text = 'Actual: Not Relevant'
cm_table.rows[1].cells[1].text = '1'
cm_table.rows[1].cells[2].text = '1'
cm_table.rows[2].cells[0].text = 'Actual: Relevant'
cm_table.rows[2].cells[1].text = '0'
cm_table.rows[2].cells[2].text = '4'
for row in cm_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
# Bold headers
for cell in cm_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for row in cm_table.rows:
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

doc.add_paragraph()
doc.add_paragraph('    - Test Accuracy: 83.33% (5 out of 6 correct)')
doc.add_paragraph('    - Overall Accuracy (full dataset): 88.89% (16 out of 18 correct)')
doc.add_paragraph('    - Precision (Relevant class): 0.80')
doc.add_paragraph('    - Recall (Relevant class): 1.00')
doc.add_paragraph('    - F1-Score (Relevant class): 0.89')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Full Dataset Predictions (Predict sheet):').bold = True

predict_data = [
    ('Software failure', 'Testing', '9.60', 'Relevant', 'Relevant', 'Correct'),
    ('Software failure', 'Design', '1.20', 'Not Relevant', 'Not Relevant', 'Correct'),
    ('Software failure', 'End of Life (EOL)', '1.20', 'Not Relevant', 'Not Relevant', 'Correct'),
    ('Software failure', 'End of Vendor Support', '0.80', 'Not Relevant', 'Not Relevant', 'Correct'),
    ('Software failure', 'Vulnerability', '1.20', 'Not Relevant', 'Not Relevant', 'Correct'),
    ('Software failure', 'Hacking/Malware', '1.60', 'Not Relevant', 'Not Relevant', 'Correct'),
    ('Hardware Failure', 'Natural/Man-made disasters', '5.00', 'Relevant', 'Relevant', 'Correct'),
    ('Hardware Failure', 'Hard drive failures', '10.00', 'Relevant', 'Relevant', 'Correct'),
    ('Hardware Failure', 'Network failures', '5.00', 'Relevant', 'Not Relevant', 'Incorrect'),
    ('Hardware Failure', 'Loss of Power', '5.00', 'Relevant', 'Relevant', 'Correct'),
    ('Users & computers', 'Passwords never expire', '4.00', 'Not Relevant', 'Relevant', 'Incorrect'),
    ('Users & computers', 'Passwords not req. to expire', '2.67', 'Not Relevant', 'Not Relevant', 'Correct'),
    ('Users & computers', 'Inactive User accounts', '6.67', 'Relevant', 'Relevant', 'Correct'),
    ('Permissions', 'Administrative permissions', '6.67', 'Relevant', 'Relevant', 'Correct'),
    ('Permissions', 'Empty security groups', '6.67', 'Relevant', 'Relevant', 'Correct'),
    ('Permissions', 'Administrative Groups', '6.67', 'Relevant', 'Relevant', 'Correct'),
    ('Data', 'Shared folders accessible', '10.00', 'Relevant', 'Relevant', 'Correct'),
    ('Data', 'Sensitive data files', '10.00', 'Relevant', 'Relevant', 'Correct'),
]

table7 = doc.add_table(rows=1, cols=6)
table7.style = 'Table Grid'
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table7.rows[0].cells
for i, h in enumerate(['Event', 'Sub-Event', 'Risk Factor', 'Actual', 'SVM Prediction', 'Result']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

for row_data in predict_data:
    row = table7.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    # Highlight incorrect rows
    if row_data[5] == 'Incorrect':
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph()
doc.add_paragraph(
    'Conclusion: The SVM model with RBF kernel successfully classified 16 out of 18 risk sub-events '
    'with 88.89% accuracy. The two misclassifications occurred at the boundary (risk factors of 4.0 and 5.0), '
    'which are close to the median threshold of 5.0. The model correctly identified all high-risk items '
    '(risk factor > 6) and all low-risk items (risk factor < 2) without errors.'
)

# ============================================================
# Note about tools
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Note:')
run.bold = True
run.underline = True

doc.add_paragraph(
    'The Vaimal Excel Add-in and MeaningCloud Add-on could not be installed due to compatibility issues '
    'with the current system environment. As an alternative, Python (with scikit-learn, pandas, and openpyxl) '
    'was used to perform the same analyses:'
)
doc.add_paragraph(
    '    - Topic extraction from unstructured email data (replacing MeaningCloud) using keyword-based '
    'NLP matching against the Risk Matrix taxonomy'
)
doc.add_paragraph(
    '    - Linear Regression (replacing Vaimal\'s regression tool) using scikit-learn\'s LinearRegression '
    'with Min-Max normalized features'
)
doc.add_paragraph(
    '    - SVM Classification (replacing Vaimal\'s SVM tool) using scikit-learn\'s SVC with RBF kernel'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Output Files:').bold = True
doc.add_paragraph('    - Results Excel: Week7-Risk_Assessment_Results.xlsx (6 sheets: Risk Assessment, '
                  'Normalized, Linear Regression, TechTrain, TechTest, Predict)')
doc.add_paragraph('    - Python Code: homework7_solution.py')
doc.add_paragraph()
doc.add_paragraph('GitHub Repository:')
doc.add_paragraph('    https://github.com/rajeshsahoo2006/CollegeProject/tree/main/Week%207%20AI%20Risk')

# Save
doc.save('/Users/sahoo/Desktop/NNDL/Code/Deep-Learning-with-TensorFlow-2-and-Keras/Week 7 AI Risk/Homework7_Submission.docx')
print("Document saved: Homework7_Submission.docx")
